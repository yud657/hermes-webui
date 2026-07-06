from __future__ import annotations

import io
import posixpath
import zipfile
import zlib
from pathlib import Path

CLAIMED_OFFICE_EXTENSIONS = frozenset({".docx", ".xlsx", ".pptx"})
CLAIMED_OFFICE_FORMATS = frozenset({"docx", "xlsx", "pptx"})
OFFICE_PREVIEW_KIND = "office"
OFFICE_RENDER_MODE = "code"
OFFICE_DEPENDENCY_HINT = (
    "Office preview is not available on this server. Install python-docx, "
    "openpyxl, and python-pptx to enable it: pip install python-docx openpyxl "
    "python-pptx"
)
OFFICE_PREVIEW_TRUNCATED_NOTICE = "[Preview truncated: Office content exceeds safe limits]"
MAX_OFFICE_PREVIEW_CHARS = 120_000
MAX_DOCX_PREVIEW_BLOCKS = 2_000
MAX_DOCX_TABLE_CELLS = 5_000
MAX_XLSX_PREVIEW_SHEETS = 20
MAX_XLSX_PREVIEW_ROWS_PER_SHEET = 500
MAX_XLSX_PREVIEW_CELLS_PER_SHEET = 5_000
MAX_PPTX_PREVIEW_SLIDES = 100
MAX_PPTX_PREVIEW_SHAPES_PER_SLIDE = 200
MAX_OFFICE_ARCHIVE_MEMBERS = 256
MAX_OFFICE_ARCHIVE_TOTAL_UNCOMPRESSED_BYTES = 8_000_000
MAX_OFFICE_ARCHIVE_MEMBER_BYTES = 4_000_000
MAX_OFFICE_ARCHIVE_MAX_COMPRESSION_RATIO = 200
MAX_DOCX_ARCHIVE_DOCUMENT_BYTES = 4_000_000
MAX_XLSX_ARCHIVE_SHARED_STRINGS_BYTES = 4_000_000
MAX_XLSX_ARCHIVE_WORKSHEET_BYTES = 2_000_000
MAX_XLSX_ARCHIVE_METADATA_BYTES = 512_000
MAX_PPTX_ARCHIVE_SLIDE_BYTES = 1_000_000
MAX_PPTX_ARCHIVE_MEDIA_BYTES = 2_000_000

_WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_DEFAULT_DOCX_SECTION_SIGNATURE = None

_DOCX_BODY_CHILDREN = {f"{_WORD_NAMESPACE}p", f"{_WORD_NAMESPACE}sectPr"}
_DOCX_PARAGRAPH_CHILDREN = {f"{_WORD_NAMESPACE}pPr", f"{_WORD_NAMESPACE}r"}
_DOCX_SAFE_PARAGRAPH_PROPERTY_CHILDREN = {f"{_WORD_NAMESPACE}pStyle"}
_DOCX_RUN_CHILDREN = {f"{_WORD_NAMESPACE}t"}


def _office_dependency_import_error() -> ImportError:
    return ImportError(OFFICE_DEPENDENCY_HINT)

def _office_archive_limit_error() -> ValueError:
    return ValueError("Office preview exceeds safe archive limits")

def _office_preview_read_error(office_format: str) -> ValueError:
    return ValueError(f"Unable to read {office_format.upper()} preview")


def _load_docx_document():
    try:
        from docx import Document as document_factory
    except ImportError as exc:  # pragma: no cover - depends on local install shape
        raise _office_dependency_import_error() from exc
    return document_factory


def _load_workbook_reader():
    try:
        from openpyxl import load_workbook as workbook_reader
    except ImportError as exc:  # pragma: no cover - depends on local install shape
        raise _office_dependency_import_error() from exc
    return workbook_reader


def _load_presentation_ctor():
    try:
        from pptx import Presentation as presentation_ctor
    except ImportError as exc:  # pragma: no cover - depends on local install shape
        raise _office_dependency_import_error() from exc
    return presentation_ctor

def _normalise_archive_member_name(name: str) -> str:
    normalized = posixpath.normpath(name.replace("\\", "/"))
    if name.startswith(("/", "\\")) or normalized in {"", ".", ".."} or normalized.startswith("../"):
        raise _office_archive_limit_error()
    return normalized

def _archive_member_byte_limit(office_format: str, member_name: str) -> int:
    if office_format == "docx":
        if member_name == "word/document.xml":
            return MAX_DOCX_ARCHIVE_DOCUMENT_BYTES
    elif office_format == "xlsx":
        if member_name == "xl/sharedStrings.xml":
            return MAX_XLSX_ARCHIVE_SHARED_STRINGS_BYTES
        if member_name.startswith("xl/worksheets/"):
            return MAX_XLSX_ARCHIVE_WORKSHEET_BYTES
        if member_name.startswith("xl/theme/") or member_name.startswith("docProps/") or member_name == "xl/workbook.xml":
            return MAX_XLSX_ARCHIVE_METADATA_BYTES
    elif office_format == "pptx":
        if member_name.startswith("ppt/slides/"):
            return MAX_PPTX_ARCHIVE_SLIDE_BYTES
        if member_name.startswith("ppt/media/"):
            return MAX_PPTX_ARCHIVE_MEDIA_BYTES
    return MAX_OFFICE_ARCHIVE_MEMBER_BYTES

def _preflight_office_archive(office_format: str, raw: bytes) -> None:
    try:
        archive = zipfile.ZipFile(io.BytesIO(raw))
    except zipfile.BadZipFile as exc:
        raise _office_preview_read_error(office_format) from exc

    with archive:
        file_infos = [info for info in archive.infolist() if not info.is_dir()]
        if len(file_infos) > MAX_OFFICE_ARCHIVE_MEMBERS:
            raise _office_archive_limit_error()

        total_uncompressed = 0
        for info in file_infos:
            member_name = _normalise_archive_member_name(info.filename)
            member_limit = _archive_member_byte_limit(office_format, member_name)
            member_size = 0
            try:
                member = archive.open(info)
            except Exception as exc:  # pragma: no cover - malformed archive path
                raise _office_preview_read_error(office_format) from exc
            with member:
                try:
                    while True:
                        chunk = member.read(64 * 1024)
                        if not chunk:
                            break
                        member_size += len(chunk)
                        total_uncompressed += len(chunk)
                        if member_size > member_limit or total_uncompressed > MAX_OFFICE_ARCHIVE_TOTAL_UNCOMPRESSED_BYTES:
                            raise _office_archive_limit_error()
                except (zipfile.BadZipFile, zlib.error, EOFError, OSError) as exc:
                    # Corrupt/truncated member data (bad CRC, short read, etc.) —
                    # mundane for partially-uploaded files, not adversarial. Raise
                    # the module's intended read-error ValueError (handled as a
                    # clean 4xx) instead of letting BadZipFile/zlib.error escape to
                    # the route's catch-all as an unhandled 500 + traceback. The
                    # limit-exceeded ValueError above is deliberately NOT caught
                    # here (it is not one of these decompression exception types).
                    raise _office_preview_read_error(office_format) from exc
            if member_size > max(info.compress_size, 1) * MAX_OFFICE_ARCHIVE_MAX_COMPRESSION_RATIO:
                raise _office_archive_limit_error()


def is_claimed_office_path(path: str | Path) -> bool:
    return Path(str(path)).suffix.lower() in CLAIMED_OFFICE_EXTENSIONS


def _office_format_for_path(path: str | Path) -> str:
    return Path(str(path)).suffix.lower().lstrip(".")


def _normalise_preview_text(value, max_chars: int | None = None) -> str:
    if value is None:
        return ""
    text = str(value)
    if max_chars is not None and max_chars >= 0 and len(text) > max_chars:
        text = text[:max_chars]
    return text.replace("\r", "\n").replace("\n", " ").strip()


def _preview_line_count(content: str) -> int:
    if not content:
        return 1
    return content.count("\n") + 1


def _finalize_preview_text(content: str, truncated: bool = False, strip_edges: bool = True) -> tuple[str, bool]:
    # docx passes strip_edges=False: leading/trailing blank paragraphs are
    # meaningful body content and the editor textarea is prefilled from this
    # text, so stripping edge whitespace would silently drop those paragraphs on
    # an unedited save (interior blanks already round-trip). xlsx/pptx keep the
    # strip — their previews are read-only and edge whitespace is just noise.
    text = (content or "").strip() if strip_edges else (content or "")
    if len(text) > MAX_OFFICE_PREVIEW_CHARS:
        text = text[:MAX_OFFICE_PREVIEW_CHARS].rstrip()
        truncated = True
    if truncated:
        text = f"{text}\n\n{OFFICE_PREVIEW_TRUNCATED_NOTICE}" if text else OFFICE_PREVIEW_TRUNCATED_NOTICE
    return text, truncated


class _PreviewBuilder:
    def __init__(self, char_limit: int | None = None) -> None:
        self._char_limit = MAX_OFFICE_PREVIEW_CHARS if char_limit is None else max(char_limit, 0)
        self._parts: list[str] = []
        self._length = 0
        self._started = False
        self.truncated = False

    @property
    def remaining_chars(self) -> int:
        return max(self._char_limit - self._length, 0)

    @property
    def started(self) -> bool:
        return self._started

    @property
    def has_content(self) -> bool:
        return bool(self._parts)

    @property
    def text(self) -> str:
        return "".join(self._parts)

    def _append_piece(self, piece: str) -> bool:
        if self.truncated:
            return False
        if not piece:
            return True
        remaining = self.remaining_chars
        if remaining <= 0:
            self.truncated = True
            return False
        if len(piece) > remaining:
            piece = piece[:remaining].rstrip()
            self.truncated = True
        if piece:
            self._parts.append(piece)
            self._length += len(piece)
        return not self.truncated

    def start_line(self) -> bool:
        if self._started:
            return self._append_piece("\n")
        self._started = True
        return True

    def start_section(self) -> bool:
        if self._started:
            return self._append_piece("\n\n")
        self._started = True
        return True

    def append_text(self, text: str) -> bool:
        if not self._started:
            self._started = True
        return self._append_piece(text)

    def finish(self, strip_edges: bool = True) -> tuple[str, bool]:
        return _finalize_preview_text(self.text, self.truncated, strip_edges=strip_edges)


def _append_normalized_preview_text(builder: _PreviewBuilder, value) -> bool:
    if builder.truncated:
        return False
    remaining = builder.remaining_chars
    if remaining <= 0:
        builder.truncated = True
        return False
    raw_text = "" if value is None else str(value)
    clipped = len(raw_text) > remaining
    if not builder.append_text(_normalise_preview_text(raw_text, remaining)):
        return False
    if clipped:
        builder.truncated = True
        return False
    return True


def _append_verbatim_preview_text(builder: _PreviewBuilder, value) -> bool:
    """Append run text WITHOUT per-node strip/space-normalization (docx runs).

    A docx paragraph is split across multiple ``<w:t>`` runs and the whitespace
    at run boundaries is significant: ``p.add_run("Hello ") + p.add_run("world")``
    stores two runs whose text is ``"Hello "`` and ``"world"``. The general
    ``_normalise_preview_text`` path ``.strip()``s each node and concatenates
    with no separator, which corrupts that to ``"Helloworld"`` — and because the
    editor textarea is prefilled from the preview text, opening + saving such a
    file (even with no edits) persisted the corruption to disk. That
    normalization is still correct for xlsx cells / pptx shapes (whole-cell
    values, not run fragments), so only the docx run path switches to verbatim
    append. Text is budget-clipped but never stripped/space-collapsed.
    """
    if builder.truncated:
        return False
    remaining = builder.remaining_chars
    if remaining <= 0:
        builder.truncated = True
        return False
    raw_text = "" if value is None else str(value)
    clipped = len(raw_text) > remaining
    if not builder.append_text(raw_text):
        return False
    if clipped:
        builder.truncated = True
        return False
    return True


def _iter_docx_text_nodes(element):
    for node in element.iter():
        if node.tag == f"{_WORD_NAMESPACE}t" and node.text:
            yield node.text


def _append_docx_element_text(builder: _PreviewBuilder, element) -> bool:
    for text in _iter_docx_text_nodes(element):
        # Verbatim (not _normalise_preview_text): preserve run-boundary
        # whitespace so multi-run paragraphs round-trip losslessly. See
        # _append_verbatim_preview_text for why docx differs from xlsx/pptx.
        if not _append_verbatim_preview_text(builder, text):
            return False
    return True


def _append_docx_cell_text(builder: _PreviewBuilder, cell_element) -> bool:
    first_paragraph = True
    for child in cell_element:
        if child.tag != f"{_WORD_NAMESPACE}p":
            continue
        if not first_paragraph and not builder.append_text("\n"):
            return False
        if not _append_docx_element_text(builder, child):
            return False
        first_paragraph = False
    return True


def _docx_preview_text(document) -> tuple[str, bool]:
    builder = _PreviewBuilder()
    body_blocks_seen = 0
    table_cells_seen = 0
    table_index = 0
    for child in document._element.body:
        if child.tag == f"{_WORD_NAMESPACE}sectPr":
            continue
        body_blocks_seen += 1
        if body_blocks_seen > MAX_DOCX_PREVIEW_BLOCKS:
            builder.truncated = True
            break
        if child.tag == f"{_WORD_NAMESPACE}p":
            if not builder.start_line() or not _append_docx_element_text(builder, child):
                break
            continue
        if child.tag != f"{_WORD_NAMESPACE}tbl":
            continue
        table_index += 1
        if not builder.start_line() or not builder.append_text(f"Table {table_index}"):
            break
        for row in child:
            if row.tag != f"{_WORD_NAMESPACE}tr":
                continue
            if not builder.start_line():
                break
            first_cell = True
            for cell in row:
                if cell.tag != f"{_WORD_NAMESPACE}tc":
                    continue
                table_cells_seen += 1
                if table_cells_seen > MAX_DOCX_TABLE_CELLS:
                    builder.truncated = True
                    break
                if not first_cell and not builder.append_text("\t"):
                    break
                if not _append_docx_cell_text(builder, cell):
                    break
                first_cell = False
            if builder.truncated:
                break
        if builder.truncated:
            break
    # strip_edges=False: preserve leading/trailing blank paragraphs so an
    # unedited open->save round-trips them (the editor prefills from this text).
    return builder.finish(strip_edges=False)


def _docx_paragraph_properties_are_safe(properties) -> bool:
    for child in properties:
        if child.tag not in _DOCX_SAFE_PARAGRAPH_PROPERTY_CHILDREN:
            return False
        if child.tag == f"{_WORD_NAMESPACE}pStyle" and child.get(f"{_WORD_NAMESPACE}val") != "Normal":
            return False
    return True


def _docx_xml_signature(element) -> tuple:
    attributes = tuple(
        sorted(
            (key, value)
            for key, value in element.attrib.items()
            if not key.rsplit("}", 1)[-1].startswith("rsid")
        )
    )
    children = tuple(_docx_xml_signature(child) for child in element)
    text = (element.text or "").strip()
    return element.tag, attributes, text, children


def _default_docx_section_signature() -> tuple:
    global _DEFAULT_DOCX_SECTION_SIGNATURE
    if _DEFAULT_DOCX_SECTION_SIGNATURE is None:
        document = _load_docx_document()()
        _DEFAULT_DOCX_SECTION_SIGNATURE = tuple(
            _docx_xml_signature(child) for child in document._element.body.sectPr
        )
    return _DEFAULT_DOCX_SECTION_SIGNATURE


def _docx_section_properties_are_safe(section_properties) -> bool:
    return tuple(_docx_xml_signature(child) for child in section_properties) == _default_docx_section_signature()


def _docx_editability(document) -> tuple[bool, str | None]:
    body = document._element.body
    for child in body:
        if child.tag not in _DOCX_BODY_CHILDREN:
            return False, "docx contains unsupported structures"
        if child.tag == f"{_WORD_NAMESPACE}sectPr" and not _docx_section_properties_are_safe(child):
            return False, "docx contains unsupported section content"
    for paragraph in document.paragraphs:
        for child in paragraph._p:
            if child.tag not in _DOCX_PARAGRAPH_CHILDREN:
                return False, "docx contains unsupported paragraph structures"
            if child.tag == f"{_WORD_NAMESPACE}pPr" and not _docx_paragraph_properties_are_safe(child):
                return False, "docx contains unsupported paragraph structures"
        for run in paragraph.runs:
            for child in run._r:
                if child.tag not in _DOCX_RUN_CHILDREN:
                    return False, "docx contains unsupported inline content"
    return True, None


def _preview_docx(raw: bytes) -> tuple[str, bool, str | None, bool]:
    _preflight_office_archive("docx", raw)
    try:
        document = _load_docx_document()(io.BytesIO(raw))
    except ImportError:
        raise
    except Exception as exc:  # pragma: no cover - library-specific failure mode
        raise ValueError("Unable to read DOCX preview") from exc
    content, truncated = _docx_preview_text(document)
    if truncated:
        return content, False, "docx preview exceeds safe limits", True
    editable, reason = _docx_editability(document)
    return content, editable, reason, truncated


def _preview_xlsx(raw: bytes) -> tuple[str, bool]:
    _preflight_office_archive("xlsx", raw)
    try:
        workbook = _load_workbook_reader()(io.BytesIO(raw), data_only=True, read_only=True)
    except ImportError:
        raise
    except Exception as exc:  # pragma: no cover - library-specific failure mode
        raise ValueError("Unable to read XLSX preview") from exc
    builder = _PreviewBuilder()
    try:
        for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
            if sheet_index > MAX_XLSX_PREVIEW_SHEETS:
                builder.truncated = True
                break
            if not builder.start_section() or not builder.append_text(f"Sheet: {sheet.title}"):
                break
            rows_seen = 0
            cells_seen = 0
            # openpyxl read-only mode reports max_row/max_column as None for any
            # workbook lacking a <dimension> record — which includes everything
            # openpyxl.Workbook(write_only=True) produces. `getattr(..., DEFAULT)`
            # does NOT help (the attribute exists; its value is None).
            #
            # When the dimension IS known, bound iter_rows to the capped extent.
            # When it's UNKNOWN, pass None so openpyxl yields each row's NATURAL
            # width — passing the cap instead would make openpyxl pad every row
            # out to `max_col` with None cells, exhausting the per-sheet cell
            # budget on row 1 and dropping the rest of the sheet. The rows_seen /
            # cells_seen counters below bound the work in the unknown case.
            sheet_max_row = getattr(sheet, "max_row", None)
            sheet_max_col = getattr(sheet, "max_column", None)
            max_row = min(sheet_max_row, MAX_XLSX_PREVIEW_ROWS_PER_SHEET) if sheet_max_row else None
            max_col = min(sheet_max_col, MAX_XLSX_PREVIEW_CELLS_PER_SHEET) if sheet_max_col else None
            for row in sheet.iter_rows(values_only=True, max_row=max_row, max_col=max_col):
                rows_seen += 1
                if rows_seen > MAX_XLSX_PREVIEW_ROWS_PER_SHEET:
                    builder.truncated = True
                    break
                row_budget = builder.remaining_chars - (1 if builder.started else 0)
                row_builder = _PreviewBuilder(row_budget)
                for value in row:
                    cells_seen += 1
                    if cells_seen > MAX_XLSX_PREVIEW_CELLS_PER_SHEET:
                        builder.truncated = True
                        break
                    if row_builder.has_content and not row_builder.append_text("\t"):
                        break
                    if not _append_normalized_preview_text(row_builder, value):
                        if not row_builder.truncated:
                            builder.truncated = True
                        break
                if builder.truncated:
                    break
                if row_builder.has_content:
                    if not builder.start_line() or not builder.append_text(row_builder.text):
                        break
                    if row_builder.truncated:
                        builder.truncated = True
                        break
            if builder.truncated:
                break
            if (getattr(sheet, "max_row", None) or rows_seen) > MAX_XLSX_PREVIEW_ROWS_PER_SHEET:
                builder.truncated = True
                break
            if (getattr(sheet, "max_column", None) or 0) > MAX_XLSX_PREVIEW_CELLS_PER_SHEET:
                builder.truncated = True
                break
    finally:
        close = getattr(workbook, "close", None)
        if callable(close):
            close()
    if not builder.has_content:
        return _finalize_preview_text("Empty workbook", builder.truncated)
    return builder.finish()


def _preview_pptx(raw: bytes) -> tuple[str, bool]:
    _preflight_office_archive("pptx", raw)
    try:
        presentation = _load_presentation_ctor()(io.BytesIO(raw))
    except ImportError:
        raise
    except Exception as exc:  # pragma: no cover - library-specific failure mode
        raise ValueError("Unable to read PPTX preview") from exc
    builder = _PreviewBuilder()
    for slide_index, slide in enumerate(presentation.slides, start=1):
        if slide_index > MAX_PPTX_PREVIEW_SLIDES:
            builder.truncated = True
            break
        if not builder.start_section() or not builder.append_text(f"Slide {slide_index}"):
            break
        shapes_seen = 0
        has_text = False
        for shape in slide.shapes:
            shapes_seen += 1
            if shapes_seen > MAX_PPTX_PREVIEW_SHAPES_PER_SLIDE:
                builder.truncated = True
                break
            text = getattr(shape, "text", "")
            if not _normalise_preview_text(text):
                continue
            if not builder.start_line() or not _append_normalized_preview_text(builder, text):
                break
            has_text = True
        if builder.truncated:
            break
        if not has_text and (not builder.start_line() or not builder.append_text("(empty slide)")):
            break
    if not builder.has_content:
        return _finalize_preview_text("Empty presentation", builder.truncated)
    return builder.finish()


def preview_office_document(path: str | Path, raw: bytes) -> dict:
    office_format = _office_format_for_path(path)
    if office_format not in CLAIMED_OFFICE_FORMATS:
        raise ValueError(f"Unsupported Office format: {path}")

    truncated = False
    if office_format == "docx":
        content, editable, reason, truncated = _preview_docx(raw)
    elif office_format == "xlsx":
        content, truncated = _preview_xlsx(raw)
        editable, reason = False, "xlsx preview is read-only in this slice"
    elif office_format == "pptx":
        content, truncated = _preview_pptx(raw)
        editable, reason = False, "pptx preview is read-only in this slice"
    else:  # pragma: no cover - exhaustive guard
        raise ValueError(f"Unsupported Office format: {path}")

    payload = {
        "path": str(path),
        "content": content,
        "size": len(raw),
        "lines": _preview_line_count(content),
        "preview_kind": OFFICE_PREVIEW_KIND,
        "office_format": office_format,
        "render_mode": OFFICE_RENDER_MODE,
        "editable": editable,
    }
    if reason:
        payload["edit_blocked_reason"] = reason
    if truncated:
        payload["truncated"] = True
    return payload


def _docx_bytes_from_text(content: str, current_bytes: bytes | None = None) -> bytes:
    """Rebuild a docx from edited plain text, preserving the original package.

    We reload the CURRENT document (styles.xml, docProps/core.xml, theme,
    settings, sectPr) and replace only the body's paragraph content, rather than
    starting from python-docx's blank template. Building from the blank template
    silently wiped author/title/custom styles on an unedited open→save — the
    same fail-closed round-trip class the sectPr and run-whitespace fixes closed.
    When current_bytes is unavailable we fall back to a fresh document.

    Content is bounded BEFORE the (quadratic — python-docx scans for sectPr on
    every add_paragraph) build loop: an editable preview is capped at
    MAX_OFFICE_PREVIEW_CHARS and MAX_DOCX_PREVIEW_BLOCKS lines, so any legitimate
    editor save fits well within these bounds. Rejecting oversized input up front
    prevents a write-surface CPU/RSS DoS (measured ~20s CPU for 50k lines).
    """
    text = str(content or "").replace("\r\n", "\n").replace("\r", "\n")
    if len(text) > MAX_OFFICE_PREVIEW_CHARS:
        raise ValueError("DOCX content exceeds the editable size limit")
    lines = text.split("\n")
    if len(lines) > MAX_DOCX_PREVIEW_BLOCKS:
        raise ValueError("DOCX content exceeds the editable paragraph limit")

    if current_bytes is not None:
        try:
            document = _load_docx_document()(io.BytesIO(current_bytes))
        except ImportError:
            raise
        except Exception:
            document = _load_docx_document()()
    else:
        document = _load_docx_document()()
    body = document._element.body
    for child in list(body):
        if child.tag != f"{_WORD_NAMESPACE}sectPr":
            body.remove(child)
    for line in lines:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def save_office_document(path: str | Path, current_bytes: bytes, content: str) -> tuple[dict, bytes]:
    office_format = _office_format_for_path(path)
    if office_format != "docx":
        raise ValueError(f"{office_format or 'office file'} is preview-only in this slice")

    current_preview = preview_office_document(path, current_bytes)
    if not current_preview.get("editable"):
        raise ValueError(current_preview.get("edit_blocked_reason") or "DOCX document is not editable")

    # Rebuild from the CURRENT package so styles/docProps/theme survive an
    # unedited round-trip; the body is still fully replaced and re-verified.
    saved_bytes = _docx_bytes_from_text(content, current_bytes)
    saved_preview = preview_office_document(path, saved_bytes)
    if not saved_preview.get("editable"):
        raise ValueError(saved_preview.get("edit_blocked_reason") or "Saved DOCX is not editable")
    return saved_preview, saved_bytes
