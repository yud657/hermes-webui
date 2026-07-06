from __future__ import annotations

import io
import json
from pathlib import Path
import shutil
import subprocess
from types import SimpleNamespace
from urllib.parse import urlparse

import pytest

# Optional Office parsers — importorskip so a lean install (parsers absent)
# skips this file cleanly rather than aborting collection suite-wide.
pytest.importorskip("docx")
pytest.importorskip("openpyxl")

from docx import Document as DocxDocument
from openpyxl import Workbook

import api.office_documents as office_documents
import api.routes as routes


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_JS = (ROOT / "static" / "workspace.js").read_text(encoding="utf-8")
NODE = shutil.which("node")


def _simple_docx_bytes(*paragraphs: str) -> bytes:
    document = DocxDocument()
    for paragraph in paragraphs or ("",):
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _simple_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "alpha"
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _extract_workspace_block(start_marker: str, end_marker: str) -> str:
    start = WORKSPACE_JS.find(start_marker)
    assert start >= 0, f"{start_marker!r} not found in static/workspace.js"
    end = WORKSPACE_JS.find(end_marker, start)
    assert end >= 0, f"{end_marker!r} not found after {start_marker!r}"
    return WORKSPACE_JS[start:end]


def _extract_workspace_function(name: str) -> str:
    marker = f"async function {name}("
    start = WORKSPACE_JS.find(marker)
    assert start >= 0, f"{name} not found in static/workspace.js"
    params_depth = 0
    body_start = -1
    for idx in range(start, len(WORKSPACE_JS)):
        char = WORKSPACE_JS[idx]
        if char == "(":
            params_depth += 1
        elif char == ")":
            params_depth -= 1
        elif char == "{" and params_depth == 0:
            body_start = idx
            break
    assert body_start >= 0, f"could not find function body for {name}"
    depth = 0
    for idx in range(body_start, len(WORKSPACE_JS)):
        char = WORKSPACE_JS[idx]
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return WORKSPACE_JS[start:idx + 1]
    raise AssertionError(f"could not find balanced function body for {name}")


def _patch_file_ops(monkeypatch, workspace: Path):
    session = SimpleNamespace(workspace=str(workspace))
    captured = {}

    def fake_j(handler, payload, status=200, extra_headers=None):
        captured["ok"] = payload
        captured["status"] = status
        return True

    def fake_bad(handler, message, status=400):
        captured["bad"] = (message, status)
        return True

    monkeypatch.setattr(routes, "get_session_for_file_ops", lambda sid: session)
    monkeypatch.setattr(routes, "j", fake_j)
    monkeypatch.setattr(routes, "bad", fake_bad)
    return captured


def test_file_read_returns_office_preview_payload_for_docx(tmp_path, monkeypatch):
    workspace = tmp_path / "ws"
    workspace.mkdir()
    (workspace / "story.docx").write_bytes(_simple_docx_bytes("alpha", "beta"))
    captured = _patch_file_ops(monkeypatch, workspace)

    routes._handle_file_read(object(), urlparse("/api/file?session_id=sid&path=story.docx"))

    payload = captured["ok"]
    assert payload["preview_kind"] == "office"
    assert payload["office_format"] == "docx"
    assert payload["render_mode"] == "code"
    assert payload["editable"] is True
    assert payload["content"] == "alpha\nbeta"


def test_file_read_returns_503_when_office_parsers_are_missing(tmp_path, monkeypatch):
    workspace = tmp_path / "ws"
    workspace.mkdir()
    (workspace / "story.docx").write_bytes(_simple_docx_bytes("alpha"))
    captured = _patch_file_ops(monkeypatch, workspace)

    def fail_read(*_args, **_kwargs):
        raise ImportError(office_documents.OFFICE_DEPENDENCY_HINT)

    monkeypatch.setattr(routes, "read_file_content", fail_read)

    routes._handle_file_read(object(), urlparse("/api/file?session_id=sid&path=story.docx"))

    assert captured["bad"] == (office_documents.OFFICE_DEPENDENCY_HINT, 503)


def test_office_save_route_accepts_safe_docx_and_rejects_preview_only_formats(tmp_path, monkeypatch):
    workspace = tmp_path / "ws"
    workspace.mkdir()
    docx_path = workspace / "story.docx"
    docx_path.write_bytes(_simple_docx_bytes("alpha", "beta"))
    xlsx_path = workspace / "budget.xlsx"
    xlsx_path.write_bytes(_simple_xlsx_bytes())

    captured = _patch_file_ops(monkeypatch, workspace)

    routes._handle_office_file_save(
        object(),
        {"session_id": "sid", "path": "story.docx", "content": "alpha\nbeta\ngamma"},
    )
    saved = captured["ok"]
    assert saved["preview_kind"] == "office"
    assert saved["office_format"] == "docx"
    assert saved["editable"] is True
    assert saved["content"] == "alpha\nbeta\ngamma"
    assert [paragraph.text for paragraph in DocxDocument(io.BytesIO(docx_path.read_bytes())).paragraphs] == [
        "alpha",
        "beta",
        "gamma",
    ]

    captured.clear()
    routes._handle_office_file_save(
        object(),
        {"session_id": "sid", "path": "budget.xlsx", "content": "ignored"},
    )
    assert captured["bad"][1] == 400
    assert "preview-only" in captured["bad"][0]


def test_office_save_route_returns_503_when_office_parsers_are_missing(tmp_path, monkeypatch):
    workspace = tmp_path / "ws"
    workspace.mkdir()
    (workspace / "story.docx").write_bytes(_simple_docx_bytes("alpha"))
    captured = _patch_file_ops(monkeypatch, workspace)

    def fail_save(*_args, **_kwargs):
        raise ImportError(office_documents.OFFICE_DEPENDENCY_HINT)

    monkeypatch.setattr(office_documents, "save_office_document", fail_save)

    routes._handle_office_file_save(
        object(),
        {"session_id": "sid", "path": "story.docx", "content": "beta"},
    )

    assert captured["bad"] == (office_documents.OFFICE_DEPENDENCY_HINT, 503)


def test_workspace_js_docx_routes_through_file_read_not_download():
    assert "'.docx'" not in WORKSPACE_JS.split("const DOWNLOAD_EXTS = new Set([", 1)[1].split("]);", 1)[0]
    assert "'.xlsx'" not in WORKSPACE_JS.split("const DOWNLOAD_EXTS = new Set([", 1)[1].split("]);", 1)[0]
    assert "'.pptx'" not in WORKSPACE_JS.split("const DOWNLOAD_EXTS = new Set([", 1)[1].split("]);", 1)[0]
    assert "'.doc'" in WORKSPACE_JS
    assert "'.xls'" in WORKSPACE_JS
    assert "'.ppt'" in WORKSPACE_JS
    assert "data.preview_kind==='office'" in WORKSPACE_JS
    assert "/api/file/office-save" in WORKSPACE_JS


def test_workspace_js_edit_button_uses_server_editable_flag():
    assert "_previewServerEditable" in WORKSPACE_JS
    assert "_previewServerEditable===false" in WORKSPACE_JS
    assert "This Office document is preview-only." in WORKSPACE_JS
    assert "_previewSaveRoute" in WORKSPACE_JS
    assert "_previewSaveRoute = data.preview_kind==='office' ? '/api/file/office-save' : '/api/file/save';" in WORKSPACE_JS


@pytest.mark.skipif(NODE is None, reason="node not on PATH")
def test_openfile_download_only_path_preserves_existing_office_save_route():
    js = (
        _extract_workspace_block("const DOWNLOAD_EXTS = new Set([", "function fileExt")
        + "function fileExt(p){ const i=p.lastIndexOf('.'); return i>=0?p.slice(i).toLowerCase():''; }\n"
        + _extract_workspace_function("openFile")
        + "\nconst S = { session: { session_id: 'sid-1' } };\n"
        + "let _previewServerEditable = true;\n"
        + "let _previewSaveRoute = '/api/file/office-save';\n"
        + "let _previewOfficeFormat = 'docx';\n"
        + "let _previewPreviewKind = 'office';\n"
        + "let downloaded = null;\n"
        + "function downloadFile(path){ downloaded = path; }\n"
        + "(async()=>{ await openFile('legacy.doc'); console.log(JSON.stringify({downloaded,_previewServerEditable,_previewSaveRoute,_previewOfficeFormat,_previewPreviewKind})); })();\n"
    )
    result = subprocess.run([NODE, "-e", js], check=True, capture_output=True, text=True, timeout=30)
    state = json.loads(result.stdout.strip().splitlines()[-1])

    assert state["downloaded"] == "legacy.doc"
    assert state["_previewSaveRoute"] == "/api/file/office-save"
    assert state["_previewServerEditable"] is True
    assert state["_previewOfficeFormat"] == "docx"
    assert state["_previewPreviewKind"] == "office"


def test_pdf_preview_path_is_unchanged():
    assert "showPreview('pdf')" in WORKSPACE_JS
    assert "_workspaceRouteForPath(path, 'raw', {inline:true})" in WORKSPACE_JS


def test_zip_and_legacy_office_formats_still_download():
    download_block = WORKSPACE_JS.split("const DOWNLOAD_EXTS = new Set([", 1)[1].split("]);", 1)[0]
    for ext in [".doc", ".xls", ".ppt", ".odt", ".ods", ".odp", ".zip", ".tar"]:
        assert ext in download_block
